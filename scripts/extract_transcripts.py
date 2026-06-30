#!/usr/bin/env python3
"""Extract all .docx/.txt transcripts in the repo into a consolidated dataset.

Produces:
  - data/transcripts.json   full per-file record (raw text + detected workflow/automation lines)
  - data/transcripts.csv    flat row-per-file summary for spreadsheet/CSV consumers
  - data/MEMORY.md          condensed index for Claude Code to load into memory
"""
import csv
import json
import re
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent.parent
DATA_DIR = ROOT / "data"
DATA_DIR.mkdir(exist_ok=True)

# Headings/lines that signal a workflow or automation step worth pulling out.
WORKFLOW_PATTERN = re.compile(
    r"\b(workflow|automation|automate|pipeline|step \d+|trigger|schedule|cron|webhook|api call|n8n|zapier|make\.com|integration)\b",
    re.IGNORECASE,
)


def read_docx(path: Path) -> str:
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def extract_workflow_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if WORKFLOW_PATTERN.search(line)]


def main() -> None:
    records = []
    files = sorted(list(ROOT.glob("*.docx")) + list(ROOT.glob("*.txt")))

    for path in files:
        try:
            text = read_docx(path) if path.suffix == ".docx" else read_txt(path)
        except Exception as exc:  # corrupt/unreadable file
            text = ""
            print(f"WARN: failed to read {path.name}: {exc}")

        workflow_lines = extract_workflow_lines(text)
        records.append(
            {
                "filename": path.name,
                "filetype": path.suffix.lstrip("."),
                "char_count": len(text),
                "word_count": len(text.split()),
                "workflow_automation_lines": workflow_lines,
                "text": text,
            }
        )
        print(f"extracted {path.name}: {len(text)} chars, {len(workflow_lines)} workflow/automation lines")

    (DATA_DIR / "transcripts.json").write_text(
        json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    csv_path = DATA_DIR / "transcripts.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(
            ["filename", "filetype", "char_count", "word_count", "workflow_automation_line_count", "workflow_automation_lines"]
        )
        for r in records:
            writer.writerow(
                [
                    r["filename"],
                    r["filetype"],
                    r["char_count"],
                    r["word_count"],
                    len(r["workflow_automation_lines"]),
                    " || ".join(r["workflow_automation_lines"]),
                ]
            )

    memory_lines = ["# Hollow Ridge Transcript Memory Index", ""]
    for r in records:
        memory_lines.append(f"## {r['filename']}")
        memory_lines.append(f"- words: {r['word_count']}, workflow/automation hits: {len(r['workflow_automation_lines'])}")
        if r["workflow_automation_lines"]:
            for line in r["workflow_automation_lines"][:10]:
                memory_lines.append(f"  - {line}")
        memory_lines.append("")
    (DATA_DIR / "MEMORY.md").write_text("\n".join(memory_lines), encoding="utf-8")

    print(f"\nDone. {len(records)} files processed -> {DATA_DIR}")


if __name__ == "__main__":
    main()
